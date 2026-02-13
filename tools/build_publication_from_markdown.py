from __future__ import annotations

import argparse
from copy import deepcopy
import re
from pathlib import Path

import latex2mathml.converter as latex2mathml
from lxml import etree
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(".")
SRC_MD = Path("paper/pgups_2026/article_mic_ieee_vak_pgups.md")
MML2OMML_CANDIDATES = [
    Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL"),
]

_MATH_XSLT: etree.XSLT | None = None


def _set_font(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _get_math_transform() -> etree.XSLT | None:
    global _MATH_XSLT
    if _MATH_XSLT is not None:
        return _MATH_XSLT
    for candidate in MML2OMML_CANDIDATES:
        if candidate.exists():
            _MATH_XSLT = etree.XSLT(etree.parse(str(candidate)))
            return _MATH_XSLT
    return None


def _latex_to_omml(latex_expr: str) -> tuple[etree._Element | None, str | None]:
    # Extract \tag{n} to avoid rendering the literal "tag" token inside equation body.
    tag_match = re.search(r"\\tag\s*\{([^}]+)\}", latex_expr)
    eq_tag = tag_match.group(1) if tag_match else None
    expr = re.sub(r"\\tag\s*\{[^}]+\}", "", latex_expr).strip()
    if not expr:
        return None, eq_tag
    transform = _get_math_transform()
    if transform is None:
        return None, eq_tag
    try:
        mathml = latex2mathml.convert(expr)
        mml_root = etree.fromstring(mathml.encode("utf-8"))
        omml_tree = transform(mml_root)
        return omml_tree.getroot(), eq_tag
    except Exception:
        return None, eq_tag


def _add_run_text(p, text: str, *, size: int, bold: bool, italic: bool) -> None:
    if not text:
        return
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)


def _add_text_with_inline_math(p, text: str, *, size: int, bold: bool, italic: bool) -> None:
    # Split by inline math fragments: $...$
    # This keeps prose in regular runs and inserts OMML for math fragments.
    cursor = 0
    for m in re.finditer(r"\$(.+?)\$", text):
        if m.start() > cursor:
            _add_run_text(p, text[cursor:m.start()], size=size, bold=bold, italic=italic)
        expr = m.group(1).strip()
        omml, _ = _latex_to_omml(expr)
        if omml is not None:
            p._p.append(deepcopy(omml))
        else:
            _add_run_text(p, m.group(0), size=size, bold=bold, italic=italic)
        cursor = m.end()
    if cursor < len(text):
        _add_run_text(p, text[cursor:], size=size, bold=bold, italic=italic)


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_page_layout(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)


def _add_page_number(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    _set_font(run, size=10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _add_par(
    doc: Document,
    text: str,
    *,
    size: int = 14,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_cm: float = 1.0,
    space_before: float = 0.0,
    space_after: float = 6.0,
    line_spacing: float = 1.0,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(first_line_cm) if first_line_cm > 0 else None
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if "$" in text:
        _add_text_with_inline_math(p, text, size=size, bold=bold, italic=italic)
    else:
        _add_run_text(p, text, size=size, bold=bold, italic=italic)


def _add_equation_block(doc: Document, latex_block: str) -> None:
    omml, eq_tag = _latex_to_omml(latex_block)
    if omml is None:
        _add_par(
            doc,
            f"$${latex_block}$$",
            size=12,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_cm=0.0,
            space_after=6.0,
        )
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(2.0)
    p.paragraph_format.space_after = Pt(6.0)
    p.paragraph_format.line_spacing = 1.0
    p._p.append(deepcopy(omml))
    if eq_tag:
        run = p.add_run(f" ({eq_tag})")
        _set_font(run, size=12)


def _clean_md_inline(text: str) -> str:
    # Keep the content and strip lightweight markdown markers.
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("*", "")
    return text.strip()


def _parse_md_table(lines: list[str], start: int) -> tuple[int, list[list[str]]]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        ln = lines[i].strip()
        if not ln.startswith("|"):
            break
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return i, rows


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    data_rows = [r for idx, r in enumerate(rows) if idx != 1]  # remove markdown separator row
    if not data_rows:
        return
    ncols = len(data_rows[0])
    t = doc.add_table(rows=1, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_idx, cell_text in enumerate(data_rows[0]):
        cell = t.rows[0].cells[c_idx]
        cell.text = _clean_md_inline(cell_text)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _set_font(run, size=12, bold=True)

    for row_data in data_rows[1:]:
        row = t.add_row().cells
        for c_idx in range(ncols):
            value = _clean_md_inline(row_data[c_idx]) if c_idx < len(row_data) else ""
            row[c_idx].text = value
            for p in row[c_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    _set_font(run, size=12)
    _add_par(doc, "", first_line_cm=0.0, space_after=2.0, line_spacing=1.0)


def _resolve_figure_path(raw: str) -> Path:
    candidate = Path(raw.strip().strip("\"'"))
    if candidate.is_absolute():
        return candidate

    relative = Path(str(candidate).replace("/", "\\"))
    probes = [
        SRC_MD.parent / relative,
        ROOT / relative,
        Path.cwd() / relative,
    ]
    for p in probes:
        if p.exists():
            return p
    return relative


def _extract_figure_path(line: str) -> Path | None:
    # Match plain path mention (legacy format).
    m = re.search(
        r"(?:`)?([A-Za-z]:\\[^`]+?\.(?:png|jpg|jpeg)|outputs\/[^`]+?\.(?:png|jpg|jpeg)|[A-Za-z0-9_./\\\\-]+?\.(?:png|jpg|jpeg))(?:`)?",
        line,
        re.IGNORECASE,
    )
    if not m:
        return None
    return _resolve_figure_path(m.group(1))


def _extract_md_image(line: str) -> tuple[str, Path] | None:
    m = re.search(r"!\[([^\]]*)\]\(([^)]+)\)", line)
    if not m:
        return None
    alt = _clean_md_inline(m.group(1))
    path = _resolve_figure_path(m.group(2))
    return alt, path


def _figure_caption(path: Path) -> str:
    key = path.stem.lower()
    mapping = {
        "fig_algorithm_block_ru": "Блок-схема алгоритма MIC AI",
        "fig_power_saving_ru": "Экономия входной мощности по сценариям",
        "fig_mech_characteristics_ru": "Сравнение механических характеристик FOC и MIC",
        "fig_power_eta_time_ru": "Сравнение потребления и КПД во времени (режим «пуск—стоп»)",
        "fig_ablation_methods_ru": "Абляционное сравнение вариантов MIC",
        "fig_ablation_methods": "Абляционное сравнение вариантов MIC",
        "fig_tradeoff_idref": "Компромисс при выборе параметра id_ref",
        "fig_timeseries_start_stop": "Временные диаграммы режима start_stop",
        "fig_timeseries_hold": "Временные диаграммы установившегося режима",
    }
    return mapping.get(key, path.stem.replace("_", " "))

def _format_fig_caption(*, raw_caption: str, fig_idx: int) -> str:
    cap = raw_caption.strip()
    if not cap:
        cap = f"Рис. {fig_idx}."
    # Normalize common caption prefixes to the journal style.
    cap = re.sub(r"^\s*Рисунок\s*", "Рис. ", cap, flags=re.IGNORECASE)
    # If caption already contains figure number, keep it as-is (author-controlled numbering).
    if re.match(r"^\s*Рис\.\s*\d+", cap, flags=re.IGNORECASE):
        out = cap
    else:
        # Strip any residual "Рис."/"Fig." prefixes without number to avoid duplication.
        cap = re.sub(r"^\s*(Рис\.|Fig\.|Figure)\s*", "", cap, flags=re.IGNORECASE).strip(" .—-")
        out = f"Рис. {fig_idx}. {cap}".strip()
    if not out.endswith("."):
        out = f"{out}."
    return out


def _pick_default_src_md() -> Path:
    primary = Path("paper/pgups_2026/article_mic_ieee_vak_pgups.md")
    if primary.exists():
        return primary
    legacy = Path("outputs/research20260212/study_final/article_mic_ieee_vak_pgups.md")
    return legacy


def _pick_default_out_docx(src_md: Path) -> Path:
    src_norm = src_md.as_posix().replace("\\", "/")
    if src_norm.endswith("paper/pgups_2026/article_mic_ieee_vak_pgups.md"):
        return src_md.parent / "СТАТЬЯ_MIC_ПГУПС_2026.docx"
    return src_md.parent / f"{src_md.stem}.docx"


def build(*, src_md: Path, out_docx: Path) -> None:
    global ROOT, SRC_MD
    ROOT = src_md.parent
    SRC_MD = src_md

    if not SRC_MD.exists():
        raise FileNotFoundError(SRC_MD)

    text = SRC_MD.read_text(encoding="utf-8").lstrip("\ufeff")
    lines = text.splitlines()

    doc = Document()
    _set_normal_style(doc)
    _set_page_layout(doc)
    _add_page_number(doc)

    i = 0
    fig_idx = 1
    seen_body_section = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        if line == "---":
            i += 1
            _add_par(doc, "", first_line_cm=0.0, space_after=4.0, line_spacing=1.0)
            continue

        if line.startswith("# "):
            _add_par(
                doc,
                _clean_md_inline(line[2:]),
                size=14,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_cm=0.0,
                space_before=4.0,
                space_after=8.0,
            )
            i += 1
            continue

        if line.startswith("## "):
            heading = _clean_md_inline(line[3:])
            # Treat only numbered sections as body ("1. ...", "2. ...").
            if re.match(r"^\d+[\.\)]\s*", heading):
                seen_body_section = True
                _add_par(
                    doc,
                    heading,
                    size=14,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_cm=0.0,
                    space_before=8.0,
                    space_after=4.0,
                )
            else:
                _add_par(
                    doc,
                    heading,
                    size=12,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_cm=0.0,
                    space_before=6.0,
                    space_after=3.0,
                )
            i += 1
            continue

        if line.startswith("### "):
            _add_par(
                doc,
                _clean_md_inline(line[4:]),
                size=13,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_cm=0.0,
                space_before=6.0,
                space_after=3.0,
            )
            i += 1
            continue

        if line.startswith("|"):
            i, table_rows = _parse_md_table(lines, i)
            _add_table(doc, table_rows)
            continue

        if line == "$$":
            j = i + 1
            eq_lines: list[str] = []
            while j < len(lines) and lines[j].strip() != "$$":
                eq_lines.append(lines[j])
                j += 1
            latex_block = "\n".join(eq_lines).strip()
            if latex_block:
                _add_equation_block(doc, latex_block)
            i = j + 1 if j < len(lines) else j
            continue

        md_figure = _extract_md_image(line)
        if md_figure is not None:
            md_caption, md_path = md_figure
            if md_path.exists():
                doc.add_picture(str(md_path), width=Cm(15.8))
                if md_caption:
                    caption = _format_fig_caption(raw_caption=md_caption, fig_idx=fig_idx)
                else:
                    caption = _format_fig_caption(raw_caption=_figure_caption(md_path), fig_idx=fig_idx)
                _add_par(
                    doc,
                    caption,
                    size=12,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_cm=0.0,
                    space_after=8.0,
                )
                fig_idx += 1
            else:
                _add_par(
                    doc,
                    f"[missing figure: {md_path}]",
                    size=10,
                    italic=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_cm=0.0,
                )
            i += 1
            continue

        if line.startswith("- "):
            fig_path = _extract_figure_path(line)
            if fig_path is not None and fig_path.exists():
                doc.add_picture(str(fig_path), width=Cm(15.8))
                caption = _format_fig_caption(raw_caption=_figure_caption(fig_path), fig_idx=fig_idx)
                _add_par(
                    doc,
                    caption,
                    size=12,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_cm=0.0,
                    space_after=8.0,
                )
                fig_idx += 1
            else:
                _add_par(
                    doc,
                    f"• {_clean_md_inline(line[2:])}",
                    first_line_cm=0.0,
                    space_after=3.0,
                )
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m_num:
            figure_candidate = _extract_figure_path(m_num.group(2))
            if figure_candidate is not None and figure_candidate.exists():
                doc.add_picture(str(figure_candidate), width=Cm(14.2))
                caption = _format_fig_caption(raw_caption=_figure_caption(figure_candidate), fig_idx=fig_idx)
                _add_par(
                    doc,
                    caption,
                    size=12,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_cm=0.0,
                    space_after=8.0,
                )
                fig_idx += 1
            else:
                _add_par(doc, f"{m_num.group(1)}. {_clean_md_inline(m_num.group(2))}", first_line_cm=0.0, space_after=3.0)
            i += 1
            continue

        # First-page metadata formatting by requirements.
        if not seen_body_section:
            cleaned = _clean_md_inline(line)
            if cleaned.startswith("УДК"):
                _add_par(doc, cleaned, size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0)
            elif cleaned.startswith("DOI:"):
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0)
            elif cleaned.startswith("Для цитирования:"):
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0)
            elif cleaned.startswith("For citation:"):
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0)
            elif cleaned.startswith("Аннотация"):
                _add_par(doc, cleaned, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0, space_before=4.0)
            elif cleaned.startswith("Summary") or cleaned.startswith("Abstract"):
                _add_par(doc, cleaned, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0, space_before=4.0)
            elif cleaned.startswith("Ключевые слова:"):
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=0.0)
            elif cleaned.startswith("Keywords:"):
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=0.0)
            elif re.match(r"^[А-ЯA-Z]\.\s*[А-ЯA-Z]\.", cleaned):
                _add_par(doc, cleaned, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0)
            elif "Петербургский государственный университет путей сообщения" in cleaned or "Petersburg State Transport University" in cleaned:
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0.0)
            else:
                _add_par(doc, cleaned, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=0.0)
            i += 1
            continue

        # Keep formulas and other plain paragraphs as-is.
        _add_par(doc, _clean_md_inline(line))
        i += 1

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    print(out_docx.resolve())


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Build PGUPS-style DOCX from Markdown + figures.")
    ap.add_argument("--src-md", type=Path, default=None, help="Path to source Markdown (default: paper/pgups_2026/...).")
    ap.add_argument("--out-docx", type=Path, default=None, help="Output DOCX path (default: рядом с исходником).")
    args = ap.parse_args()

    src_md = args.src_md or _pick_default_src_md()
    out_docx = args.out_docx or _pick_default_out_docx(src_md)
    build(src_md=src_md, out_docx=out_docx)
